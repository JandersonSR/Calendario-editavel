# app.py
import streamlit as st
from datetime import datetime, date
import calendar
from io import BytesIO
import pandas as pd
import holidays

# docx
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# reportlab
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet

# excel
from openpyxl import Workbook
from openpyxl.styles import PatternFill, Font, Alignment, Border, Side

st.set_page_config(page_title="Calendário Interativo → Escala", layout="wide")
st.title("📅 Calendário Interativo → Gera Escala (clique nos dias)")

# ---------------- Sidebar: configurações ----------------
st.sidebar.header("Configurações")
ano = st.sidebar.number_input("Ano", min_value=1900, max_value=2100, value=datetime.now().year)
mes = st.sidebar.selectbox("Mês", list(calendar.month_name)[1:], index=datetime.now().month - 1)
usar_feriados_nacionais = st.sidebar.checkbox("Incluir feriados nacionais (Brasil)", value=True)
feriados_custom = st.sidebar.text_area("Feriados adicionais (DD/MM por linha)", help="Ex: 25/12", value="")

# layout options / colors
cor_cabecalho = st.sidebar.color_picker("Cor do cabeçalho (para export)", "#D9EDF7")
cor_amarelo = st.sidebar.color_picker("Cor destaque amarelo (para export)", "#FFFF66")
cor_verde = st.sidebar.color_picker("Cor destaque verde (para export)", "#C6E0B4")

# title and blocks to produce
titulo = st.sidebar.text_input("Título do documento:", value=f"ESCALA - {mes}/{ano}")

dias_possiveis = ["Domingo","Segunda-feira","Terça-feira","Quarta-feira","Quinta-feira","Sexta-feira","Sábado"]
blocos_padrao = st.sidebar.multiselect("Quais blocos aparecerão na escala:",
                                       options=dias_possiveis,
                                       default=["Domingo","Segunda-feira","Sexta-feira","Sábado"])


# ---------------- prepare holidays ----------------
def parse_custom_holidays(text, year):
    s = set()
    for line in text.splitlines():
        line = line.strip()
        if not line:
            continue
        try:
            d = datetime.strptime(line + f"/{year}", "%d/%m/%Y").date()
            s.add(d)
        except Exception:
            st.sidebar.warning(f"Formato inválido para feriado: '{line}'. Use DD/MM")
    return s

feriados_adicionais = parse_custom_holidays(feriados_custom, ano)
feriados_br = holidays.BR(years=[ano]) if usar_feriados_nacionais else {}


# ---------------- session state ----------------
if "date_names" not in st.session_state:
    st.session_state.date_names = {}

if "selected_date" not in st.session_state:
    st.session_state.selected_date = None


# ---------------- calendar rendering ----------------
st.markdown("### Calendário do mês")
mes_num = list(calendar.month_name).index(mes)
cal = calendar.Calendar(firstweekday=6)  # Sunday-first
month_dates = [d for d in cal.itermonthdates(ano, mes_num)]

st.markdown(f"#### {mes} / {ano}")

weekdays = ["DOM", "SEG", "TER", "QUA", "QUI", "SEX", "SÁB"]
cols = st.columns(7)
for i, wd in enumerate(weekdays):
    cols[i].markdown(f"**{wd}**")

weeks = [month_dates[i:i+7] for i in range(0, len(month_dates), 7)]

for week in weeks:
    cols = st.columns(7)
    for i, day in enumerate(week):
        col = cols[i]
        if day.month != mes_num:
            col.write("")
            continue

        day_key = day.isoformat()
        name_assigned = st.session_state.date_names.get(day_key, "")
        is_holiday = (day in feriados_adicionais) or (day in feriados_br)

        label = f"**{day.day}**"
        if name_assigned:
            label += f"\n\n{name_assigned}"
        if is_holiday:
            label += "\n\n**FERIADO**"

        if col.button(label, key=f"btn_{day_key}"):
            st.session_state.selected_date = day_key


# ---------------- selected date editor ----------------
st.markdown("---")
sel = st.session_state.selected_date
if sel:
    sel_date = datetime.fromisoformat(sel).date()
    st.markdown(f"### Editando: **{sel_date.strftime('%d/%m/%Y (%A)')}**")
    current = st.session_state.date_names.get(sel, "")

    with st.form(key=f"form_{sel}"):
        name = st.text_input("Nome para este dia (deixe vazio para remover):", value=current)
        save_btn = st.form_submit_button("Salvar")
        if save_btn:
            if name.strip():
                st.session_state.date_names[sel] = name.strip()
            else:
                st.session_state.date_names.pop(sel, None)
            st.rerun()

    if st.button("Limpar este dia"):
        st.session_state.date_names.pop(sel, None)
        st.rerun()


# ---------------- summary table ----------------
st.markdown("### Dias preenchidos")
if st.session_state.date_names:
    df_assigned = pd.DataFrame([
        {"data_fmt": datetime.fromisoformat(k).strftime("%d/%m/%Y"),
         "weekday": datetime.fromisoformat(k).strftime("%A"),
         "nome": v}
        for k, v in st.session_state.date_names.items()
    ])
    st.dataframe(df_assigned)
else:
    st.info("Nenhum dia preenchido ainda.")


# ---------------- mapping to blocks ----------------
weekday_map = {
    "Segunda-feira": 0,
    "Terça-feira": 1,
    "Quarta-feira": 2,
    "Quinta-feira": 3,
    "Sexta-feira": 4,
    "Sábado": 5,
    "Domingo": 6
}

estrutura = {}
for bloco in blocos_padrao:
    wd = weekday_map[bloco]
    estrutura[bloco] = [d for d in month_dates if (d.month == mes_num and d.weekday() == wd)]

blocos_filled = {}
for bloco, dates in estrutura.items():
    datas_fmt, nomes = [], []
    for i in range(5):
        if i < len(dates):
            d = dates[i]
            key = d.isoformat()
            nome = st.session_state.date_names.get(key, "")
            is_hol = (d in feriados_br) or (d in feriados_adicionais)
            datas_fmt.append(d.strftime("%d/%m/%Y") + (" - FERIADO" if is_hol else ""))
            nomes.append(nome)
        else:
            datas_fmt.append("-")
            nomes.append("")

    blocos_filled[bloco] = {
        "datas": datas_fmt,
        "nomes": nomes,
        "local": "",
        "horario": ""
    }


# ---------------- inputs local + horário ----------------
st.markdown("---")
st.markdown("### Local e Horário por bloco")
for bloco in blocos_filled:
    l = st.text_input(f"Local - {bloco}", key=f"local_{bloco}")
    blocos_filled[bloco]["local"] = l

    h = st.text_input(f"Horário - {bloco}", key=f"horario_{bloco}")
    blocos_filled[bloco]["horario"] = h

observacoes = st.text_area("Observações:", height=120)


# ---------------- DOCX ----------------
def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color.replace("#",""))
    tcPr.append(shd)

def gerar_docx(titulo, blocos_filled, cor_cabecalho, cor_amarelo, cor_verde, observacoes):
    doc = Document()

    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run(titulo.upper())
    run.bold = True
    run.font.size = Pt(14)
    doc.add_paragraph()

    for bloco, content in blocos_filled.items():
        datas = content["datas"]
        nomes = content["nomes"]
        local = content["local"]
        horario = content["horario"]

        t1 = doc.add_table(rows=1, cols=1)
        c = t1.rows[0].cells[0]
        set_cell_bg(c, cor_amarelo)
        doc.add_paragraph()

        tbl = doc.add_table(rows=4, cols=7)
        tbl.autofit = False
        widths = [Cm(2.0), Cm(2.2)] + [Cm(3.2)]*5
        for j,w in enumerate(widths):
            for cell in tbl.columns[j].cells:
                cell.width = w

        first = tbl.cell(0,0)
        for k in range(1,7):
            first = first.merge(tbl.cell(0,k))
        p_title = tbl.rows[0].cells[0].paragraphs[0]
        p_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = p_title.add_run(bloco)
        run.bold = True
        run.font.size = Pt(12)
        set_cell_bg(tbl.rows[0].cells[0], cor_cabecalho)

        # headers
        tbl.rows[1].cells[0].text = "Local"
        set_cell_bg(tbl.rows[1].cells[0], cor_amarelo)
        tbl.rows[1].cells[1].text = "Horário"
        set_cell_bg(tbl.rows[1].cells[1], cor_cabecalho)

        for i in range(5):
            tbl.rows[1].cells[2+i].text = f"{i+1}º {bloco}"
            set_cell_bg(tbl.rows[1].cells[2+i], cor_cabecalho)

        # dates
        tbl.rows[2].cells[0].text = local or ""
        tbl.rows[2].cells[1].text = horario or ""
        for i in range(5):
            tbl.rows[2].cells[2+i].text = datas[i]

        # names
        for i in range(5):
            tbl.rows[3].cells[2+i].text = nomes[i]

        doc.add_paragraph()

    doc.add_paragraph("Observações:")
    doc.add_paragraph(observacoes)
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio


# ---------------- PDF ----------------
def gerar_pdf(titulo, blocos_filled, cor_cabecalho, cor_amarelo, cor_verde, observacoes):
    buf = BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=landscape(A4))
    styles = getSampleStyleSheet()
    elements = []

    elements.append(Paragraph(titulo, styles['Title']))
    elements.append(Spacer(1,8))

    for bloco, content in blocos_filled.items():
        datas = content["datas"]
        nomes = content["nomes"]
        local = content["local"]
        horario = content["horario"]

        sp = Table([[""]], colWidths=[720/7])
        sp.setStyle(TableStyle([('BACKGROUND',(0,0),(-1,-1), colors.HexColor(cor_amarelo))]))
        elements.append(sp)
        elements.append(Spacer(1,6))

        elements.append(Paragraph(f"<b>{bloco}</b>", styles['Heading3']))
        elements.append(Spacer(1,6))

        header = ["Local", "Horário"] + [f"{i+1}º {bloco}" for i in range(5)]
        row1 = [local, horario] + datas
        row2 = ["", ""] + nomes

        tbl = Table([header, row1, row2], colWidths=[60,80] + [90]*5)
        sty = TableStyle([
            ('GRID',(0,0),(-1,-1),0.8,colors.black),
            ('BACKGROUND',(0,0),(-1,0), colors.HexColor(cor_cabecalho)),
            ('ALIGN',(0,0),(-1,-1),'CENTER'),
        ])
        tbl.setStyle(sty)
        elements.append(tbl)
        elements.append(Spacer(1,12))

    elements.append(Paragraph("<b>Observações:</b>", styles['Normal']))
    elements.append(Paragraph(observacoes, styles['Normal']))

    doc.build(elements)
    buf.seek(0)
    return buf


# ---------------- EXCEL ----------------
def gerar_excel(titulo, blocos_filled, cor_cabecalho, cor_amarelo, cor_verde, observacoes):

    def hex_to_fill(h):
        return PatternFill(start_color=h.replace("#",""), end_color=h.replace("#",""), fill_type="solid")

    wb = Workbook()
    ws = wb.active
    ws.title = "Escala"

    border = Border(
        left=Side(style="thin"), right=Side(style="thin"),
        top=Side(style="thin"), bottom=Side(style="thin")
    )

    ws.merge_cells("A1:G1")
    title_cell = ws["A1"]
    title_cell.value = titulo.upper()
    title_cell.font = Font(bold=True, size=14)
    title_cell.alignment = Alignment(horizontal="center")

    row = 3

    for bloco, content in blocos_filled.items():
        datas = content["datas"]
        nomes = content["nomes"]
        local = content["local"]
        horario = content["horario"]

        ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=7)
        c = ws.cell(row,1)
        c.fill = hex_to_fill(cor_amarelo)
        row += 1

        ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=7)
        c = ws.cell(row,1, bloco)
        c.font = Font(bold=True, size=12)
        c.alignment = Alignment(horizontal="center")
        c.fill = hex_to_fill(cor_cabecalho)
        row += 1

        headers = ["Local","Horário"] + [f"{i+1}º {bloco}" for i in range(5)]
        for col, h in enumerate(headers, start=1):
            cell = ws.cell(row, col, h)
            cell.fill = hex_to_fill(cor_cabecalho)
            cell.font = Font(bold=True)
            cell.alignment = Alignment(horizontal="center")
            cell.border = border
        row += 1

        values = [local, horario] + datas
        for col, v in enumerate(values, start=1):
            cell = ws.cell(row, col, v)
            cell.alignment = Alignment(horizontal="center")
            cell.border = border
        row += 1

        values = ["",""] + nomes
        for col, v in enumerate(values, start=1):
            cell = ws.cell(row, col, v)
            cell.alignment = Alignment(horizontal="center")
            cell.border = border

        row += 2

    ws2 = wb.create_sheet("Observações")
    ws2["A1"] = "Observações:"
    ws2["A1"].font = Font(bold=True)
    ws2["A2"] = observacoes

    for col in "ABCDEFG":
        ws.column_dimensions[col].width = 18

    bio = BytesIO()
    wb.save(bio)
    bio.seek(0)
    return bio


# ---------------- Buttons ----------------
col1, col2, col3, col4 = st.columns(4)

with col1:
    if st.button("Gerar DOCX"):
        docx_bytes = gerar_docx(titulo, blocos_filled, cor_cabecalho, cor_amarelo, cor_verde, observacoes)
        st.download_button("Baixar DOCX", data=docx_bytes.getvalue(),
                           file_name=f"escala_{mes}_{ano}.docx")

with col2:
    if st.button("Gerar PDF"):
        pdf_bytes = gerar_pdf(titulo, blocos_filled, cor_cabecalho, cor_amarelo, cor_verde, observacoes)
        st.download_button("Baixar PDF", data=pdf_bytes.getvalue(),
                           file_name=f"escala_{mes}_{ano}.pdf")

with col3:
    rows = []
    for k,v in st.session_state.date_names.items():
        d = datetime.fromisoformat(k).date()
        rows.append({"data": d.strftime("%d/%m/%Y"), "dia": d.strftime("%A"), "nome": v})
    dfout = pd.DataFrame(rows)
    st.download_button("Baixar CSV", data=dfout.to_csv(index=False),
                       file_name=f"preenchimentos_{mes}_{ano}.csv")

with col4:
    if st.button("Gerar Excel"):
        excel_bytes = gerar_excel(titulo, blocos_filled, cor_cabecalho, cor_amarelo, cor_verde, observacoes)
        st.download_button("Baixar Excel", data=excel_bytes.getvalue(),
                           file_name=f"escala_{mes}_{ano}.xlsx",
                           mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

st.success("Pronto! O sistema está funcionando com DOCX, PDF e EXCEL.")
